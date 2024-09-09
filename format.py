# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
# from docx.enum.style import WD_STYLE_TYPE
# # Function to format heading
# def format_heading(paragraph, level=1):
#     style_name = f'Heading {level}'
#     paragraph.style = style_name
#     if level == 1:
#         paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     return paragraph

# # Function to format bullet points
# def format_bullets(paragraph):
#     paragraph.style = 'List Bullet'
#     return paragraph

# # Function to format normal text
# def format_normal_text(paragraph):
#     paragraph.style = 'Normal'
#     return paragraph
# # Function to format bullet points
# def format_bullets(paragraph):
#     doc = paragraph.part.document
#     if 'List Bullet' not in doc.styles:
#         style = doc.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
#         style.base_style = doc.styles['Normal']
#         style.paragraph_format.left_indent = Pt(18)
#         style.paragraph_format.first_line_indent = Pt(-18)
#     paragraph.style = 'List Bullet'
#     return paragraph
# # Function to check and format document
# def format_document(doc_path, output_path):
#     doc = Document(doc_path)

#     for para in doc.paragraphs:
#         text = para.text.strip()

#         # Checking and applying styles
#         if text.lower().startswith('heading 1:'):
#             para = format_heading(para, level=1)
#             para.text = text.replace('Heading 1:', '').strip()
        
#         elif text.lower().startswith('heading 2:'):
#             para = format_heading(para, level=2)
#             para.text = text.replace('Heading 2:', '').strip()
        
#         elif para.text.startswith('-') or para.text.startswith('•'):
#             para = format_bullets(para)

#         else:
#             para = format_normal_text(para)

#     # Save the formatted document
#     doc.save(output_path)
#     print(f'Document formatted and saved as {output_path}')

# # Usage
# input_file = 'D:\GSPs\IBIB Final.docx'  # Your input Word document
# output_file = 'ibib_final.docx'  # Output Word document with proper formatting
# format_document(input_file, output_file)









from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

# Function to format heading
def format_heading(paragraph, level=1):
    style_name = f'Heading {level}'
    paragraph.style = style_name
    if level == 1:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return paragraph

# Function to format bullet points
def format_bullets(paragraph):
    paragraph.style = 'List Bullet'
    return paragraph

# Function to format normal text
def format_normal_text(paragraph):
    paragraph.style = 'Normal'
    return paragraph

def format_bullets(paragraph):
    doc = paragraph.part.document
    if 'List Bullet' not in doc.styles:
        style = doc.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles['Normal']
        style.paragraph_format.left_indent = Pt(18)
        style.paragraph_format.first_line_indent = Pt(-18)
    paragraph.style = 'List Bullet'
    return paragraph

# Function to create the 'Caption' style if it doesn't exist
def add_caption_style(doc):
    styles = doc.styles
    if 'Caption' not in [s.name for s in styles]:
        # Create new style for caption
        style = styles.add_style('Caption', 1)
        style.font.size = Pt(12)
        style.font.bold = True
    return doc

# Function to label figures (images)
def label_figure(paragraph, figure_num):
    paragraph.style = 'Caption'
    paragraph.text = f'Figure {figure_num}: ' + paragraph.text.strip()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return paragraph

# Function to check and format document
def format_document(doc_path, output_path):
    doc = Document(doc_path)
    doc = add_caption_style(doc)  # Ensure the caption style exists
    figure_count = 1

    for para in doc.paragraphs:
        text = para.text.strip()

        # Check if the paragraph contains an image (figure)
        if 'rId' in para._p.xml:  # Detect image reference inside paragraph
            print(f'Figure detected in paragraph: {text}')
            para = label_figure(para, figure_num=figure_count)
            figure_count += 1

        # Checking and applying styles for headings
        elif text.lower().startswith('heading 1:'):
            print(f'Heading 1 detected: {text}')
            para = format_heading(para, level=1)
            para.text = text.replace('Heading 1:', '').strip()
        
        elif text.lower().startswith('heading 2:'):
            print(f'Heading 2 detected: {text}')
            para = format_heading(para, level=2)
            para.text = text.replace('Heading 2:', '').strip()
        
        # Checking for bullet points
        elif para.text.startswith('-') or para.text.startswith('•'):
            print(f'Bullet point detected: {text}')
            para = format_bullets(para)

        # Normal text formatting
        else:
            print(f'Normal text detected: {text}')
            para = format_normal_text(para)


    # Save the formatted document
    doc.save(output_path)
    print(f'Document formatted and saved as {output_path}')

# Usage
input_file = 'D:\GSPs\IBIB Final.docx'  # Your input Word document
output_file = '1.docx'  # Output Word document with proper formatting
format_document(input_file, output_file)










#fig

# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.oxml import OxmlElement

# # Function to format heading
# def format_heading(paragraph, level=1):
#     style_name = f'Heading {level}'
#     paragraph.style = style_name
#     if level == 1:
#         paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     return paragraph

# # Function to format bullet points
# def format_bullets(paragraph):
#     paragraph.style = 'List Bullet'
#     return paragraph

# # Function to format normal text
# def format_normal_text(paragraph):
#     paragraph.style = 'Normal'
#     return paragraph

# # Function to add figure captions
# def add_figure_caption(doc, figure_index, description="Figure Description"):
#     caption = f'Figure {figure_index}: {description}'
#     caption_para = doc.add_paragraph(caption)
#     caption_para.style = 'Caption'
#     caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# # Function to check and format document
# def format_document(doc_path, output_path):
#     doc = Document(doc_path)
#     figure_index = 1

#     for para in doc.paragraphs:
#         text = para.text.strip()

#         # Checking and applying styles
#         if text.lower().startswith('heading 1:'):
#             para = format_heading(para, level=1)
#             para.text = text.replace('Heading 1:', '').strip()

#         elif text.lower().startswith('heading 2:'):
#             para = format_heading(para, level=2)
#             para.text = text.replace('Heading 2:', '').strip()

#         elif para.text.startswith('-') or para.text.startswith('•'):
#             para = format_bullets(para)

#         else:
#             para = format_normal_text(para)

#     # Handle figures (images)
#     for rel in doc.part.rels.values():
#         if "image" in rel.target_ref:
#             # Add figure caption after the paragraph where the figure was found
#             add_figure_caption(doc, figure_index, description="Sample figure caption")
#             figure_index += 1

#     # Save the formatted document
#     doc.save(output_path)
#     print(f'Document formatted with
